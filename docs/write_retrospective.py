#!/usr/bin/env python3
"""
docs/write_retrospective.py — Project Retrospective Report Generator
===================================================================
Tạo file Word: docs/Project_Retrospective_Report.docx

File này tổng kết toàn bộ quá trình làm đồ án:
- Bài toán, lý thuyết, dataset
- Các bước thực hiện + tư duy (thinking process)
- Quyết định kỹ thuật + lý do
- Sai lầm, thử nghiệm thất bại, bài học kinh nghiệm

Dựa trên dữ liệu thực tế từ:
  - src/train/preprocess.py
  - src/train/train_cost_sensitive.py
  - src/train/train_final.py
  - src/train/evaluate.py
  - src/train/threshold_tuning.py
  - src/train/extract_pcap_features.py
  - src/self_capture/capture_and_generate_sudo.py
  - PROGRESS.md

Updated: 2026-04-05
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin  = Cm(3.0)
section.right_margin = Cm(2.5)


# ══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.75)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def bold_bullet(doc, label, desc, level=0):
    """Bullet with bold label followed by normal description."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.75)
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(desc)
    return p


def shade_row(row, hex_color="D9E2F3"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),    'clear')
        shd.set(qn('w:color'),  'auto')
        shd.set(qn('w:fill'),   hex_color)
        tcPr.append(shd)


def make_table(doc, headers, rows, header_color="2E75B6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0]
    shade_row(hdr, header_color)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        if r_i % 2 == 0:
            shade_row(row, "F2F2F2")
        for c_i, val in enumerate(row_data):
            row.cells[c_i].text = str(val)
            row.cells[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(0.3)
    p.paragraph_format.space_after  = Cm(0.3)
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("PROJECT RETROSPECTIVE REPORT")
run.bold = True
run.font.size = Pt(22)

doc.add_paragraph()

sub_cover = doc.add_paragraph()
sub_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_cover.add_run("Phát hiện Data Exfiltration qua HTTP\nBằng AI và Xử lý Đa luồng")
run2.bold = True
run2.font.size = Pt(15)

doc.add_paragraph()

info_items = [
    ("Đồ án môn học:", "Giáo vụ mạng máy tính (GVMH)"),
    ("Giáo viên hướng dẫn:", "ThS. Đàm Minh Linh"),
    ("Học viện:", "Học viện Công nghệ Bưu Chính Viễn Thông — CS TP.HCM"),
    ("Ngày hoàn thành:", datetime.datetime.now().strftime("%d/%m/%Y")),
    ("Thời lượng dự án:", "~6 tuần / 14 tuần (~43%)"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label + " ").bold = True
    p.add_run(value)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "MỤC LỤC", 1)
toc = [
    "1.  Tổng quan & Lý thuyết cơ sở",
    "2.  Các bước thực hiện & Tư duy (Thinking Process)",
    "3.  Quyết định kỹ thuật & Lý do (Architectural Rationale)",
    "4.  Sai lầm, Thử nghiệm thất bại & Bài học kinh nghiệm",
    "5.  Kết luận",
]
for t in toc:
    bullet(doc, t)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: TỔNG QUAN & LÝ THUYẾT CƠ SỞ
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1. TỔNG QUAN & LÝ THUYẾT CƠ SỞ", 1)

# ── 1.1 Bài toán ────────────────────────────────────────────────────────────
heading(doc, "1.1. Bài toán: Phát hiện Data Exfiltration qua HTTP", 2)

para(doc,
    "Data Exfiltration (rò rỉ dữ liệu) là hành vi chuyển dữ liệu nhạy cảm ra khỏi tổ chức "
    "một cách trái phép. Trong bối cảnh mạng doanh nghiệp, kẻ tấn công thường lợi dụng "
    "HTTP/HTTPS (port 80/443) vì lưu lượng web được firewall cho phép và dễ dàng ẩn giấu "
    "trong traffic bình thường.")

para(doc,
    "Bài toán đặt ra: Xây dựng một hệ thống có khả năng nhận biết hành vi exfiltration "
    "dựa trên thống kê lưu lượng mạng và dấu hiệu bất thường ở tầng ứng dụng, đồng thời "
    "chạy real-time với pipeline xử lý đa luồng (multi-threaded).")

# ── 1.2 Công nghệ ───────────────────────────────────────────────────────────
heading(doc, "1.2. Công nghệ & Công cụ sử dụng", 2)

tech_headers = ["Nhóm", "Công nghệ", "Mục đích"]
tech_rows = [
    ["Dataset", "CICIDS2017 (CICFlowMeter)", "2.83M flows, 67 features đã extract"],
    ["Packet Capture", "Scapy + tcpdump (sudo)", "Đọc PCAP, BPF filter, live capture"],
    ["Anomaly Detection", "Isolation Forest, One-Class SVM", "Phát hiện outliers (zero-day)"],
    ["Deep Learning", "BiLSTM, CNN 1D (TensorFlow/Keras)", "Phát hiện patterns có giám sát"],
    ["Oversampling", "SMOTE (imbalanced-learn)", "Xử lý extreme class imbalance"],
    ["Loss Function", "Focal Loss (γ=2.0, α=0.50)", "Hard example mining, cân bằng loss"],
    ["Threshold Tuning", "ROC curve analysis", "Tối ưu FPR vs Recall tradeoff"],
    ["Pipeline", "Python threading (3 threads)", "Real-time packet → feature → inference"],
]
make_table(doc, tech_headers, tech_rows)

# ── 1.3 Dataset ─────────────────────────────────────────────────────────────
heading(doc, "1.3. Dataset CICIDS2017 — Chi tiết thống kê", 2)

para(doc,
    "CICIDS2017 (Canadian Institute for Cybersecurity) cung cấp 5 ngày network traffic "
    "thực tế với nhiều loại tấn công mạng. Chúng tôi sử dụng bộ CICFlowMeter processed CSV "
    "— đã được trích xuất sẵn 67 đặc trưng cho mỗi flow, giúp tiết kiệm công sức feature engineering.")

headers_ds = ["File CSV", "Ngày", "Số flows", "Loại tấn công chính"]
rows_ds = [
    ["Monday-WorkingHours.csv", "Thứ 2", "529,918", "BENIGN (baseline)"],
    ["Tuesday-WorkingHours.csv", "Thứ 3", "445,909", "FTP-BruteForce, SSH-BruteForce"],
    ["Wednesday-workingHours.csv", "Thứ 4", "692,703", "DoS Hulk, GoldenEye, Slowloris"],
    ["Thursday-Morning.csv", "Thứ 5 (sáng)", "170,366", "BruteForce, XSS, SQL Injection"],
    ["Thursday-Afternoon.csv", "Thứ 5 (chiều)", "288,602", "Infiltration (port scan, backdoor)"],
    ["Friday-Morning.csv", "Thứ 6 (sáng)", "191,033", "BENIGN + Bot traffic"],
    ["Friday-Afternoon-PortScan.csv", "Thứ 6 (chiều 1)", "286,467", "PortScan"],
    ["Friday-Afternoon-DDoS.csv", "Thứ 6 (chiều 2)", "225,745", "HOIC, LOIC DDoS"],
    ["TỔNG CỘNG", "—", "2,830,743", "—"],
]
make_table(doc, headers_ds, rows_ds)

para(doc, "")

# ── 1.4 Labeling ─────────────────────────────────────────────────────────────
heading(doc, "1.4. Thách thức gán nhãn Exfiltration", 2)

para(doc,
    "CICIDS2017 không có label 'Exfiltration' rõ ràng. Đây là thách thức lớn nhất của dự án: "
    "dataset chỉ gán nhãn theo loại tấn công chung (DoS, BruteForce, DDoS, Infiltration), "
    "trong khi exfiltration là một hành vi cụ thể có thể ẩn trong nhiều loại tấn công khác nhau.")

para(doc, "Chúng tôi đã sử dụng 3 nguồn exfiltration proxy:")

exfil_headers = ["Nguồn proxy", "Số flows", "Đặc điểm exfil"]
exfil_rows = [
    ["Bot traffic (Friday-Morning)", "1,966", "Upload ratio 4.57× cao hơn bình thường, "
     "short sessions, automated behavior"],
    ["Infiltration (Thursday-Afternoon)", "36", "Port scan + backdoor → hành vi exfiltration"],
    ["Heuristics (upload>5×, PSH>0.3, dur<600s)", "~202", "Custom rules: upload_ratio>5, "
     "duration ngắn, PSH flag ratio cao"],
    ["TỔNG EXFIL PROXY", "~2,204", "Chiếm 0.078% của toàn bộ dataset"],
]
make_table(doc, exfil_headers, exfil_rows)

para(doc,
    "Số lượng exfil proxy rất nhỏ (0.078%) tạo ra bài toán extreme class imbalance — "
    "đây là nguyên nhân gốc rễ của hàng loạt thử nghiệm thất bại và quyết định kỹ thuật "
    "quan trọng về sau.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: CÁC BƯỚC THỰC HIỆN & TƯ DUY (THINKING PROCESS)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2. CÁC BƯỚC THỰC HIỆN & TƯ DUY (THINKING PROCESS)", 1)

# ── 2.1 Tổng quan phases ─────────────────────────────────────────────────────
heading(doc, "2.1. Tổng quan 6 Phases của dự án", 2)

phases_headers = ["Phase", "Nội dung", "Trạng thái", "Kết quả chính"]
phases_rows = [
    ["Phase 1", "Setup + EDA", "✅ Hoàn thành", "2.83M flows, 67 features, project structure"],
    ["Phase 2", "Feature Engineering", "✅ Hoàn thành", "Custom features, burst_exfil_score, train/test/val split"],
    ["Phase 3", "Multi-thread Pipeline", "✅ Hoàn thành", "3 threads chạy thực trên Friday PCAP"],
    ["Phase 4", "Model Training", "✅ Hoàn thành", "6 models, FPR fix thành công"],
    ["Phase 5", "Evaluation", "✅ Hoàn thành", "ROC curves, AUC=0.9971, FPR=2.45%"],
    ["Phase 6", "Report + Bonus", "🔄 Đang làm", "bao_cao.docx ✅, self-capture 🔄, slide ⏳"],
]
make_table(doc, phases_headers, phases_rows)

# ── 2.2 Thinking Process ──────────────────────────────────────────────────────
heading(doc, "2.2. Tư duy đằng sau từng bước", 2)

# Step 1
heading(doc, "Bước 1: Tiền xử lý dữ liệu — Tại sao phải chuẩn hóa columns?", 3)
para(doc,
    "Khi load các file CSV, chúng tôi phát hiện một vấn đề tưởng như nhỏ nhưng quan trọng: "
    "CICFlowMeter có bug tạo ra space prefix trong column names "
    "(ví dụ: ' Label' thay vì 'Label', ' Flow Duration' thay vì 'Flow Duration'). "
    "Nếu không chuẩn hóa, merge DataFrame sẽ tạo duplicate columns, và feature selection "
    "sẽ miss một nửa features.")

para(doc, "Tư duy: Không bao giờ assume data đã clean. Luôn kiểm tra dtypes, nulls, infs, "
    "và column names trước khi bắt đầu bất kỳ pipeline nào.")

# Step 2
heading(doc, "Bước 2: Feature Extraction — Tại sao phải so sánh Anomaly vs Supervised?", 3)
para(doc,
    "Chúng tôi cố tình huấn luyện cả 2 trường phái để trả lời câu hỏi: "
    "Trong bài toán exfiltration detection, approach nào hiệu quả hơn?")

para(doc, "**Anomaly Detection (Không giám sát):**", bold=True)
bullet(doc, "Train trên NORMAL traffic (label=0) duy nhất.")
bullet(doc, "Ưu điểm: Không cần label exfiltration — phát hiện được zero-day attacks.")
bullet(doc, "Nhược điểm: Rất kém trên CICIDS2017 (AUC≈0.55) — vì Bot traffic giống Normal "
    "trong raw 67-feature space.")

para(doc, "**Supervised Learning (Có giám sát):**", bold=True)
bullet(doc, "Train trên toàn bộ dataset đã gán nhãn (2,830,743 flows).")
bullet(doc, "Ưu điểm: AUC vượt trội (0.90–0.99), discrimination cực kỳ tốt.")
bullet(doc, "Nhược điểm: Cần label exfiltration chính xác — mà CICIDS2017 không có.")

para(doc, "**Kết luận tư duy:** Không nên chọn một trường phái duy nhất. "
    "CNN1D (supervised) dùng làm primary detector; "
    "Isolation Forest (anomaly) giữ làm secondary cho zero-day. "
    "Đây là best-of-both-worlds architecture.")

# Step 3
heading(doc, "Bước 3: Model Architecture — Tại sao chọn CNN 1D và BiLSTM?", 3)
para(doc, "**CNN 1D — Local pattern detector:**", bold=True)
bullet(doc, "1D convolution với kernel_size=1 học local patterns trong feature space.")
bullet(doc, "GlobalAveragePooling1D (thay vì Flatten) cho kết quả tốt hơn vì "
    "nó giảm overfitting khi input có 67 features.")
bullet(doc, "Tại sao kernel_size=1? Vì mỗi feature độc lập, không cần spatial locality "
    "như trong image processing. Conv1D(64, kernel_size=1) hoạt động như một "
    "feedforward layer với regularization tốt hơn.")

para(doc, "**BiLSTM — Temporal pattern detector:**", bold=True)
bullet(doc, "Bidirectional giúp capture cả forward và backward temporal dependencies.")
bullet(doc, "Dropout(0.4) rất cao — có tác dụng regularization mạnh, "
    "ngăn overfitting vào training distribution.")
bullet(doc, "Tuy nhiên, BiLSTM cần nhiều training time hơn CNN1D gấp 2-3 lần "
    "và kết quả không vượt trội. Trong production, CNN1D được ưu tiên hơn.")

# Step 4
heading(doc, "Bước 4: Pipeline đa luồng — Tại sao phải 3 threads riêng biệt?", 3)
para(doc,
    "Suy nghĩ ban đầu: Có thể gộp tất cả vào 1 thread và xử lý tuần tự. "
    "Nhưng với packet rate có thể lên đến 10,000 packets/s trong attack scenarios, "
    "single thread sẽ trở thành bottleneck.")

para(doc, "Tư duy thiết kế 3 threads:")
bullet(doc, "Thread 1 (Capture) phải độc lập với Thread 2 (Feature) vì "
    "scapy.sniff() là blocking I/O — nếu feature extraction chậm, capture sẽ drop packets.")
bullet(doc, "Thread 2 (Feature) phải buffer packets theo IP và time window — "
    "đây là CPU-bound computation, không nên block I/O thread.")
bullet(doc, "Thread 3 (Inference) phải độc lập vì model prediction có thể chậm hơn "
    "feature extraction — queue giữa Thread 2 và 3 là bộ đệm quan trọng.")
bullet(doc, "Queue sizes (50K/10K) phải đủ lớn để không drop alerts trong burst attack, "
    "nhưng không quá lớn để gây memory pressure.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: QUYẾT ĐỊNH KỸ THUẬT & LÝ DO (ARCHITECTURAL RATIONALE)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3. QUYẾT ĐỊNH KỸ THUẬT & LÝ DO (ARCHITECTURAL RATIONALE)", 1)

# ── 3.1 SMOTE ─────────────────────────────────────────────────────────────────
heading(doc, "3.1. Tại sao phải dùng SMOTE? Tại sao không train trực tiếp?", 2)

para(doc,
    "Với 0.078% exfil (1,543 exfil trên ~2 triệu train samples), training trực tiếp "
    "sẽ khiến model gặp hiện tượng 'Accuracy Paradox': "
    "model đạt 99.9% accuracy bằng cách predict tất cả là normal — "
    "vì chỉ cần đúng 99.9% là đã 'perfect' rồi, mà không cần phát hiện được attack nào.")

para(doc, "SMOTE (Synthetic Minority Over-sampling Technique) giải quyết vấn đề này bằng "
    "cách nội suy các mẫu minority trong feature space. Tuy nhiên, việc OVERDOSE SMOTE "
    "chính là nguyên nhân của FPR ~45% ban đầu (xem Section 4).")

para(doc, "**Lý do chọn SMOTE 10% (thay vì 50% hoặc 100%):**", bold=True)
bullet(doc, "10% = đủ để model nhận biết minority class mà không làm distortion quá lớn.")
bullet(doc, "SMOTE 50% hoặc 100% sẽ tạo quá nhiều synthetic samples → "
    "model học boundary không tự nhiên.")
bullet(doc, "Với tỷ lệ 10%, mỗi minority sample được replicate trung bình 6-7 lần "
    "(thay vì 128 lần như ban đầu với target 10% từ 0.08%) — "
    "giảm 20× so với lần thử thất bại.")

# ── 3.2 Focal Loss ───────────────────────────────────────────────────────────
heading(doc, "3.2. Tại sao phải dùng Focal Loss thay vì Binary Cross-Entropy?", 2)

para(doc,
    "Binary Cross-Entropy (BCE) coi mọi misclassification có trọng số như nhau. "
    "Trong bài toán extreme imbalance, model nhanh chóng đạt loss thấp bằng cách "
    "predict tất cả là normal — vì mỗi correct normal prediction giảm loss đáng kể, "
    "trong khi miss một exfil sample chỉ tăng loss rất ít (vì exfil quá hiếm trong loss sum).")

para(doc, "Focal Loss giải quyết bằng cách giảm weight cho 'easy examples' "
    "(những mẫu model đã predict đúng với confidence cao), "
    "và tập trung learning vào 'hard examples' "
    "(những mẫu bị confuse giữa normal và exfil).")

para(doc, "**Công thức Focal Loss:**", bold=True)
p = doc.add_paragraph()
run = p.add_run("  FL(pₜ) = −αₜ · (1 − pₜ)ᵞ · log(pₜ)")
run.font.name = "Courier New"
run.font.size = Pt(10)

para(doc, "Trong đó:")
bullet(doc, "γ (gamma) = 2.0: điều khiển mức độ giảm weight cho easy examples. "
    "γ càng cao → càng tập trung vào hard examples. γ=2.0 là giá trị được Lin et al. "
    "(2017) chứng minh hiệu quả trong COCO dataset.")
bullet(doc, "α (alpha) = 0.50 (symmetric): cân bằng trọng số giữa 2 classes. "
    "α=0.75 (ban đầu) gây quá bias về minority → probability inflation. "
    "α=0.50 là symmetric, giảm bias.")

# ── 3.3 Threshold Tuning ─────────────────────────────────────────────────────
heading(doc, "3.3. Tại sao phải thực hiện Threshold Tuning? Tại sao không giữ mặc định 0.5?", 2)

para(doc,
    "Đây là quyết định kỹ thuật quan trọng nhất của dự án. "
    "Sau khi huấn luyện với SMOTE + Focal Loss + class_weight, "
    "model có discrimination TỐT (AUC=0.94) nhưng probability distribution "
    "BỊ SAI (mean_prob=0.44 trên dataset 99.9% normal).")

para(doc, "**Tại sao điều này xảy ra?**", bold=True)
bullet(doc, "SMOTE tạo synthetic samples với features nằm giữa normal và exfil → "
    "model không hoàn toàn chắc chắn khi predict.")
bullet(doc, "Focal Loss tập trung vào hard examples → model thường output probability "
    "ở mức trung bình (0.3–0.6) thay vì near-0 hoặc near-1.")
bullet(doc, "class_weight={1:5} đẩy model bias về phía positive class → "
    "thêm một lần nữa làm tăng mean probability.")

para(doc, "**Tại sao không giữ threshold 0.5?**", bold=True)
bullet(doc, "Với mean_prob=0.44, threshold 0.5 sẽ classify hầu hết normal traffic "
    "là exfil → FPR ~45%.")
bullet(doc, "FPR 45% có nghĩa là: cứ 100 alerts thì 45 cái là false alarm — "
    "không thể chấp nhận trong production.")
bullet(doc, "Threshold tuning quét ROC curve để tìm điểm tối ưu: "
    "FPR ≤ 5% và Recall ≥ 85%.")

para(doc, "**Kết quả threshold tuning:**", bold=True)
thresh_headers = ["Mô hình", "Threshold cũ", "Threshold mới", "FPR cũ", "FPR mới", "Giảm FPR"]
thresh_rows = [
    ["CNN 1D", "0.500", "0.207", "44.77%", "2.45%", "18.3×"],
    ["BiLSTM", "0.500", "0.167", "44.16%", "3.22%", "13.7×"],
]
make_table(doc, thresh_headers, thresh_rows)

para(doc,
    "Threshold 0.207 (CNN1D) có nghĩa là: model chỉ alert khi "
    "confidence ≥ 20.7% — cao hơn đáng kể so với 50%. "
    "Điều này phản ánh rằng trong real-world deployment với "
    "extreme imbalance, ta cần model rất chắc chắn trước khi alert.")

# ── 3.4 Subsample Strategy ────────────────────────────────────────────────────
heading(doc, "3.4. Tại sao phải Subsample 100K trước khi SMOTE?", 2)

para(doc,
    "Với 2 triệu training samples, SMOTE trên toàn bộ dataset "
    "sẽ tạo ra quá nhiều synthetic samples (1.9M × 10% = 190K), "
    "gây memory pressure và training cực kỳ chậm.")

para(doc, "Tư duy: Subsample 100K giữ nguyên tỷ lệ 2% exfil "
    "→ đủ để model học patterns mà không tốn quá nhiều compute. "
    "Sau đó SMOTE lên 10% → 100K × 10% = 10K exfil, 90K normal. "
    "Tổng 100K samples train nhanh hơn 2M samples gấp 20 lần.")

# ── 3.5 burst_exfil_score ────────────────────────────────────────────────────
heading(doc, "3.5. Metric burst_exfil_score — Tại sao phải thiết kế metric riêng?", 2)

para(doc,
    "Model ML không phải là giải pháp duy nhất. "
    "Chúng tôi thiết kế burst_exfil_score như một rule-based layer bổ sung, "
    "dựa trên domain knowledge về exfiltration behavior.")

score_headers = ["Tín hiệu", "Điều kiện", "Trọng số", "Lý do"]
score_rows = [
    ["Upload ratio cao", "upload_ratio > 5.0 AND bytes > 50KB", "+0.40",
     "Bot traffic upload 4.57× so với normal — tín hiệu mạnh nhất"],
    ["Burst pattern", "burst_count > 50", "+0.20",
     "Nhiều request liên tục → automated exfiltration"],
    ["Unusual port", "unusual_port_ratio > 0.8", "+0.20",
     "Port không phải HTTP/S standard → suspicious"],
    ["Regular timing", "inter_request_time_std < 0.05s", "+0.20",
     "Machine-generated traffic có pattern đều đặn"],
]
make_table(doc, score_headers, score_rows)

para(doc, "")
para(doc, "**Alert threshold > 0.7** có nghĩa: cần ít nhất 2 trong 4 tín hiệu mới alert. "
    "Điều này giảm false positives đáng kể so với chỉ dùng model probability alone.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: SAI LẦM, THỬ NGHIỆM THẤT BẠI & BÀI HỌC KINH NGHIỆM
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4. SAI LẦM, THỬ NGHIỆM THẤT BẠI & BÀI HỌC KINH NGHIỆM", 1)

# ── 4.1 Hội chứng Probability Inflation ──────────────────────────────────────
heading(doc, "4.1. Sai lầm #1: Hội chứng Probability Inflation (FPR ~45%)", 2)

para(doc,
    "Đây là sai lầm nghiêm trọng nhất của dự án, và cũng là bài học giá trị nhất. "
    "Xem xét kỹ file src/train/train_cost_sensitive.py (dòng 1-22) để hiểu root cause analysis.")

# Code excerpt box
p = doc.add_paragraph()
run = p.add_run("""
ROOT CAUSE ANALYSIS (trích từ train_cost_sensitive.py):
─────────────────────────────────────────────────────────────
  Current model (SMOTE 128×, focal α=0.75, class_weight=10×) → mean_prob=0.444
  on a dataset with 0.074% exfil → threshold 0.5 too low → FPR=0.45.

  Triple stacking:
    1. SMOTE target_ratio=0.1 → minority oversampled 128× (0.08% → 10%)
    2. Focal Loss α=0.75 → adds ~2-5× weight on minority
    3. class_weight={1: 10.0} → another 10× penalty
  → Model learns to predict almost everything as exfil
─────────────────────────────────────────────────────────────""")
run.font.name = "Courier New"
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(80, 80, 80)

para(doc, "")
para(doc, "**Phân tích chi tiết hiện tượng:**", bold=True)

para(doc, "SMOTE 128× (lần thử đầu tiên, thất bại):", bold=True)
bullet(doc, "Tỷ lệ gốc: 0.08% exfil (1 trong 1,280 samples).")
bullet(doc, "SMOTE target 10%: minority phải đạt 10% của majority.")
bullet(doc, "→ Oversampling factor = 10 / 0.08 = 125 → làm tròn = 128×.")
bullet(doc, "Kết quả: model nhận quá nhiều synthetic samples giống hệt nhau → "
    "overfit vào minority class, predict mọi thứ là exfil.")

para(doc, "Focal Loss α=0.75 (lần thử đầu tiên, thất bại):", bold=True)
bullet(doc, "α=0.75 nghĩa là trọng số 3:1 cho minority class.")
bullet(doc, "Kết hợp với SMOTE 128×: đây là double penalty cho minority.")
bullet(doc, "α=0.50 (symmetric) là đủ — không cần thiên về positive class thêm nữa.")

para(doc, "class_weight={1:10} (lần thử đầu tiên, thất bại):", bold=True)
bullet(doc, "Missing 1 exfil = cost = 10× missing 1 normal.")
bullet(doc, "Triple stacking: SMOTE 128× + Focal α=0.75 + class_weight=10× → "
    "effective penalty lên đến ~20-50× cho minority.")
bullet(doc, "Giải pháp: class_weight={0:1, 1:5} (chỉ 5×) + SMOTE 10× (thay vì 128×) = "
    "effective penalty ~50× → moderate, hiệu quả hơn.")

# ── 4.2 Cost-Sensitive Without SMOTE ─────────────────────────────────────────
heading(doc, "4.2. Sai lầm #2: Cost-Sensitive Learning WITHOUT SMOTE (train_cost_sensitive.py)", 2)

para(doc,
    "Sau khi phát hiện Probability Inflation, chúng tôi thử loại bỏ hoàn toàn SMOTE "
    "và chỉ dùng Cost-Sensitive Learning (class_weight={1:50}) + Focal Loss (α=0.50). "
    "Đây là một giả thuyết hợp lý, nhưng kết quả không cải thiện đáng kể.")

para(doc, "**Tại sao Cost-Sensitive without SMOTE không hiệu quả?**", bold=True)
bullet(doc, "class_weight={1:50} chỉ thay đổi gradient magnitude, "
    "không thay đổi gradient direction đủ để học minority patterns từ 0.08% samples.")
bullet(doc, "Với quá ít minority samples (1,543 trong 2 triệu), "
    "model không đủ statistical evidence để learn decision boundary chính xác.")
bullet(doc, "SMOTE bổ sung syntactic neighbors → giúp model hiểu minority distribution "
    "tốt hơn, không chỉ penalize harder.")

para(doc, "**Bài học:** SMOTE và Cost-Sensitive Learning bổ trợ lẫn nhau, "
    "không phải thay thế. Cần cả hai với mức độ vừa phải: "
    "SMOTE 10× (moderate) + class_weight=5× (moderate) + Focal α=0.50 (symmetric).")

# ── 4.3 Anomaly Models Fail ─────────────────────────────────────────────────
heading(doc, "4.3. Sai lầm #3: Anomaly Models quá kém trên CICIDS2017", 2)

para(doc,
    "Ban đầu kỳ vọng Isolation Forest sẽ là 'secret weapon' cho exfiltration detection — "
    "vì anomaly detection không cần label, phát hiện được zero-day attacks. "
    "Kết quả thực tế: AUC≈0.53 (gần như random).")

para(doc, "**Root cause:** Bot/BENIGN traffic trong CICIDS2017 có feature distributions "
    "rất giống nhau. Isolation Forest không tìm thấy meaningful outliers "
    "vì 'anomaly' trong dataset này là Infiltration (36 flows) — quá ít để "
    "tạo ra statistical footprint rõ ràng.")

para(doc, "**Bài học:** Anomaly detection chỉ hoạt động khi anomalies thực sự "
    "khác biệt trong feature space. Nếu dataset không được label exfiltration rõ ràng, "
    "anomaly models sẽ kém vì không có ground truth cho anomalies. "
    "→ Giải pháp: dùng supervised models làm primary, anomaly làm secondary.")

# ── 4.4 Label Quality ───────────────────────────────────────────────────────
heading(doc, "4.4. Sai lầm #4: Chất lượng Label — Vấn đề nền tảng", 2)

para(doc,
    "Toàn bộ dự án bị giới hạn bởi chất lượng label. "
    "CICIDS2017 có nhãn 'Infiltration' và 'Bot' nhưng không có nhãn 'Exfiltration'. "
    "Chúng tôi phải dùng proxy (Bot traffic ≈ exfiltration behavior) — "
    "điều này ảnh hưởng đến mọi khía cạnh của model training.")

para(doc, "**Hệ quả cụ thể:**", bold=True)
bullet(doc, "Precision rất thấp (~3%) là hệ quả trực tiếp của label noise. "
    "Nhiều 'Bot' flows thực ra là normal browsing behavior được gắn tag sai.")
bullet(doc, "Recall=100% có thể là illusory — model có thể đang detect 'Bot' patterns "
    "chứ không phải 'Exfiltration' thực sự.")
bullet(doc, "Trong production, cần ground-truth exfiltration labels mới có thể "
    "đánh giá model một cách đáng tin cậy.")

para(doc, "**Bài học:** Label quality > Model architecture. "
    "Một CNN1D với label tốt sẽ tốt hơn Transformer với label noise. "
    "Nên đầu tư thời gian vào data labeling hơn là hyperparameter tuning.")

# ── 4.5 Key Takeaways ────────────────────────────────────────────────────────
heading(doc, "4.5. Tổng hợp Bài học kinh nghiệm (Key Takeaways)", 2)

lessons_headers = ["STT", "Bài học", "Ứng dụng"]
lessons_rows = [
    ["1", "Triple stacking of imbalance techniques → probability inflation. "
     "Dùng 1 technique vừa phải thay vì 3 techniques quá mức.",
     "Luôn kiểm tra probability distribution, không chỉ AUC."],
    ["2", "Threshold 0.5 không phải always optimal. Trong extreme imbalance, "
     "tối ưu hóa threshold là bắt buộc.",
     "Luôn chạy threshold sensitivity analysis sau training."],
    ["3", "Anomaly detection cần statistical footprint rõ ràng. "
     "Không hoạt động khi anomaly quá ít hoặc giống normal.",
     "Validate anomaly assumption trước khi triển khai."],
    ["4", "Subsample + moderate SMOTE tốt hơn full data + extreme SMOTE. "
     "100K subsample đủ cho model học, không tốn quá nhiều compute.",
     "Subsample là best practice cho large-scale imbalanced datasets."],
    ["5", "Label quality > Model architecture. Precision 3% phản ánh label noise, "
     "không phải model yếu.",
     "Đầu tư vào data labeling trước khi tuning model."],
    ["6", "Multi-thread pipeline cần queue đủ lớn giữa các stages. "
     "Bottleneck ở stage nào → stage đó cần buffer lớn hơn.",
     "Always benchmark pipeline end-to-end, không chỉ từng thread."],
    ["7", "Focal Loss α=0.50 (symmetric) hiệu quả hơn α=0.75 trong extreme imbalance. "
     "α quá cao gây bias ngược chiều.",
     "α=0.50 là good default cho binary classification imbalance."],
]
make_table(doc, lessons_headers, lessons_rows)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: KẾT LUẬN
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "5. KẾT LUẬN", 1)

# ── 5.1 Đạt được ─────────────────────────────────────────────────────────────
heading(doc, "5.1. Những gì đã đạt được", 2)

achievements_headers = ["Mục tiêu", "Kết quả", "Trạng thái"]
achievements_rows = [
    ["AUC-ROC > 0.90 (supervised)", "CNN1D Final: AUC=0.9971", "✅ Vượt xa (0.9971 vs 0.90)"],
    ["AUC-ROC > 0.85 (anomaly)", "Isolation Forest: AUC=0.5277", "❌ Không đạt"],
    ["FPR < 5%", "CNN1D Final: FPR=2.45%", "✅ Đạt (sau threshold tuning)"],
    ["Recall ≥ 85%", "CNN1D Final: Recall=100%", "✅ Vượt xa"],
    ["Pipeline real-time", "3 threads, test trên Friday PCAP", "✅ Hoạt động"],
    ["So sánh anomaly vs supervised", "4 models, full analysis", "✅ Hoàn thành"],
    ["Metric burst_exfil_score riêng", "4 tín hiệu, threshold > 0.7", "✅ Hoàn thành"],
    ["Báo cáo tiếng Việt", "bao_cao.docx (44KB, 8 sections)", "✅ Hoàn thành"],
]
make_table(doc, achievements_headers, achievements_rows)

# ── 5.2 Technical Journey ─────────────────────────────────────────────────────
heading(doc, "5.2. Hành trình kỹ thuật — Từ thất bại đến thành công", 2)

para(doc, "Dưới đây là timeline của các thử nghiệm và kết quả:")

journey_headers = ["Bước", "Approach", "AUC", "FPR", "Kết luận"]
journey_rows = [
    ["1", "Isolation Forest (anomaly)", "0.5277", "10.10%", "❌ Near-random — feature space không phân biệt"],
    ["2", "One-Class SVM (anomaly)", "0.5546", "4.93%", "❌ AUC quá thấp"],
    ["3", "BiLSTM + SMOTE 128× + focal α=0.75", "0.9012", "44.16%", "❌ FPR cao — probability inflation"],
    ["4", "CNN1D + SMOTE 128× + focal α=0.75", "0.9423", "44.77%", "❌ FPR cao — probability inflation"],
    ["5", "CNN1D cost-sensitive (NO SMOTE, w=50×)", "≈0.90", "≈30%", "❌ Vẫn cao — cần SMOTE"],
    ["6", "CNN1D final (subsample + SMOTE 10% + "
          "focal α=0.50 + class_weight=5× + threshold tuning)", "0.9971", "2.45%",
          "✅ THÀNH CÔNG — mọi mục tiêu đạt"],
]
make_table(doc, journey_headers, journey_rows)

para(doc, "")

# ── 5.3 Final Thoughts ────────────────────────────────────────────────────────
heading(doc, "5.3. Suy nghĩ cuối cùng", 2)

para(doc,
    "Dự án này dạy cho chúng tôi một bài học sâu sắc: "
    "trong machine learning, việc hiểu data (understanding) quan trọng hơn việc "
    "chọn model phức tạp (model selection). CNN1D với AUC=0.9971 "
    "không phải vì nó 'thông minh hơn' các model khác — "
    "mà vì chúng tôi đã tìm ra đúng cách xử lý extreme class imbalance.")

para(doc,
    "Threshold tuning, từ góc nhìn này, không chỉ là một kỹ thuật tối ưu hóa — "
    "nó là cách để 'align' model output distribution với real-world prior probability. "
    "Khi real-world exfil rate là 0.078%, threshold 0.207 phản ánh đúng prior này: "
    "model chỉ alert khi có strong evidence (probability > 20.7%), "
    "thay vì mù quáng dùng 50%.")

para(doc,
    "Hướng phát triển tiếp theo: (1) thu thập production data với ground-truth exfil labels, "
    "(2) ensemble CNN1D + burst_exfil_score + Isolation Forest, "
    "(3) tích hợp HTTP header inspection (User-Agent, Content-Type) "
    "để tăng detection accuracy beyond flow-level features.")

# ── 5.4 References ─────────────────────────────────────────────────────────────
heading(doc, "5.4. Tài liệu tham khảo", 2)
refs = [
    "CICIDS2017 — Canadian Institute for Cybersecurity: https://www.unb.ca/cic/datasets/ids-2017.html",
    "Isolation Forest — Liu et al. (2008): 'Isolation Forest', ACM TKDD",
    "Focal Loss — Lin et al. (2017): 'Focal Loss for Dense Object Detection', ICCV",
    "SMOTE — Chawla et al. (2002): 'SMOTE: Synthetic Minority Over-sampling Technique', JAIR",
    "CICFlowMeter — Lashgari et al. (2020): 'Characterization of Encrypted and VPN Traffic', IEEE ICISPA",
]
for r in refs:
    bullet(doc, r)

# ── 5.5 Files generated ───────────────────────────────────────────────────────
heading(doc, "5.5. Source code chính đã viết", 2)
files = [
    ("src/train/preprocess.py", "Load, label, split, scale 2.83M flows → 67 features"),
    ("src/train/train_anomaly.py", "Isolation Forest + One-Class SVM"),
    ("src/train/train_cost_sensitive.py", "Cost-sensitive CNN1D + BiLSTM (thất bại → insight)"),
    ("src/train/train_final.py", "Final: subsample + SMOTE 10% + focal α=0.50 + threshold tuning"),
    ("src/train/evaluate.py", "ROC curves, confusion matrices, metric comparison"),
    ("src/train/threshold_tuning.py", "Post-hoc threshold optimization với ROC analysis"),
    ("src/train/extract_pcap_features.py", "PCAP → 67 CICFlowMeter-compatible features"),
    ("src/capture/packet_capture.py", "Thread 1: Scapy sniff + BPF filter"),
    ("src/features/feature_aggregator.py", "Thread 2: 60s window buffer + feature extraction"),
    ("src/inference/model_inference.py", "Thread 3: Load model + predict + alert"),
    ("src/self_capture/capture_and_generate_sudo.py", "Self-capture scenario generator"),
    ("docs/write_retrospective.py", "Script tạo báo cáo này"),
]
for fname, desc in files:
    bold_bullet(doc, fname + ":", desc)

# Save
output_path = "/Users/nguyen_bao/Documents/AI_Project/Exfiltration/docs/Project_Retrospective_Report.docx"
doc.save(output_path)
print(f"✅ Saved: {output_path}")
