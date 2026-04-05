#!/usr/bin/env python3
"""
Generate bao_cao.docx from project results.
Updated: 2026-04-05 — FPR Fix with threshold tuning.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin  = Cm(3.0)
section.right_margin = Cm(2.5)

# ── Styles helpers ──────────────────────────────────────────────────────────────
def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.75)
    p.add_run(text)
    return p

def shade_row(row, hex_color="D9E2F3"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def make_table(doc, headers, rows, header_color="2E75B6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0]
    shade_row(hdr, "2E75B6")
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        if r_i % 2 == 0:
            shade_row(row, "F2F2F2")
        for c_i, val in enumerate(row_data):
            row.cells[c_i].text = str(val)
            row.cells[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # spacing
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("PHÁT HIỆN DATA EXFILTRATION QUA HTTP\nBẰNG AI VÀ XỬ LÝ ĐA LUỒNG")
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Đồ án môn học — GVMH")
sub_run.bold = True
sub_run.font.size = Pt(13)

doc.add_paragraph()

info_items = [
    ("Giáo viên hướng dẫn:", "ThS. Đàm Minh Linh"),
    ("Học viện:", "Học viện Công nghệ Bưu Chính Viễn Thông — CS TP.HCM"),
    ("Ngày nộp:", datetime.datetime.now().strftime("%d/%m/%Y")),
    ("Môn học:", "Giáo vụ mạng máy tính (GVMH)"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{label} ").bold = True
    p.add_run(value)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# MỤC LỤC
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "MỤC LỤC", 1)
toc_items = [
    "1. Tổng quan đề tài",
    "2. Tổng quan Dataset",
    "3. Trích xuất đặc trưng (Feature Engineering)",
    "4. Pipeline đa luồng",
    "5. Huấn luyện mô hình",
    "6. Đánh giá kết quả",
    "7. So sánh Anomaly-based vs Supervised",
    "8. Kết luận và đề xuất",
]
for item in toc_items:
    add_bullet(doc, item)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. TỔNG QUAN ĐỀ TÀI
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "1. TỔNG QUAN ĐỀ TÀI", 1)

set_heading(doc, "1.1. Giới thiệu bài toán", 2)
add_para(doc,
    "Data Exfiltration (rò rỉ dữ liệu) là việc chuyển dữ liệu nhạy cảm ra khỏi tổ chức một "
    "cách trái phép. Kẻ tấn công thường lợi dụng HTTP/HTTPS (port 80/443) vì lưu lượng web "
    "được cho phép qua firewall và dễ ẩn giấu trong traffic bình thường.")

set_heading(doc, "1.2. Mục tiêu", 2)
bullets = [
    "Nhận biết hành vi exfiltration dựa trên thống kê lưu lượng HTTP/HTTPS.",
    "Xây dựng pipeline xử lý đa luồng (multi-threading) để đáp ứng thời gian thực.",
    "Huấn luyện và so sánh 4 mô hình ML: Isolation Forest, One-Class SVM, BiLSTM, CNN 1D.",
    "Đề xuất metric burst_exfil_score kết hợp nhiều tín hiệu bất thường.",
]
for b in bullets:
    add_bullet(doc, b)

set_heading(doc, "1.3. Kiến trúc hệ thống", 2)
add_para(doc, "Hệ thống gồm 3 threads hoạt động song song:")
threads = [
    ("Thread 1 — Packet Capture:", "Đọc gói tin từ PCAP hoặc network interface → đẩy vào packet_queue."),
    ("Thread 2 — Feature Aggregation:", "Gom gói tin theo IP nguồn vào cửa sổ 60 giây → trích đặc trưng → đẩy vào feature_queue."),
    ("Thread 3 — Inference + Alert:", "Load model → dự đoán → tính burst_exfil_score → alert nếu phát hiện exfiltration."),
]
for label, desc in threads:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(desc)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. TỔNG QUAN DATASET
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "2. TỔNG QUAN DATASET", 1)

set_heading(doc, "2.1. Nguồn dữ liệu", 2)
add_para(doc,
    "Dataset chính: CICIDS2017 (Canadian Institute for Cybersecurity) — "
    "5 ngày network traffic với nhiều loại tấn công. "
    "Chúng tôi sử dụng bộ CICFlowMeter processed CSV (đã trích xuất 67 đặc trưng mỗi flow).")

set_heading(doc, "2.2. Thống kê Dataset", 2)
headers = ["File CSV", "Ngày", "Số flows", "Nhãn chính"]
rows = [
    ["Monday-WorkingHours.csv", "Thứ 2", "529,918", "BENIGN (baseline)"],
    ["Tuesday-WorkingHours.csv", "Thứ 3", "445,909", "FTP-BruteForce, SSH-BruteForce"],
    ["Wednesday-workingHours.csv", "Thứ 4", "692,703", "DoS Hulk, GoldenEye, Slowloris"],
    ["Thursday-Morning.csv", "Thứ 5 (sáng)", "170,366", "BruteForce, XSS, SQL Injection"],
    ["Thursday-Afternoon.csv", "Thứ 5 (chiều)", "288,602", "Infiltration (port scan, backdoor)"],
    ["Friday-Morning.csv", "Thứ 6 (sáng)", "191,033", "BENIGN + Bot"],
    ["Friday-Afternoon-PortScan.csv", "Thứ 6 (chiều 1)", "286,467", "PortScan"],
    ["Friday-Afternoon-DDoS.csv", "Thứ 6 (chiều 2)", "225,745", "HOIC, LOIC DDoS"],
    ["TỔNG CỘNG", "—", "2,830,743", "—"],
]
make_table(doc, headers, rows)

doc.add_paragraph()
set_heading(doc, "2.3. Gán nhãn Exfiltration", 2)
add_para(doc,
    "CICIDS2017 không có label exfiltration rõ ràng. Chúng tôi sử dụng hai nguồn exfil proxy:")
proxy = [
    ("Bot traffic (Friday-Morning):", "1,966 flows — có upload ratio 4.57x cao hơn bình thường, "
     "đặc trưng cho automated data upload."),
    ("Infiltration (Thursday-Afternoon):", "36 flows — port scan + backdoor → hành vi exfiltration."),
    ("Heuristics:", "Flows có upload_ratio > 100, burst_count > 50, unusual_port_ratio > 0.8 → gán nhãn exfil."),
]
for label, desc in proxy:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(desc)

add_para(doc, "Tổng exfiltration proxy: ~2,204 flows trên 2,830,743 total (0.078%) — extreme imbalance.")
add_para(doc, "Train/Test/Val split: 70/15/15, stratified by label.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "3. TRÍCH XUẤT ĐẶC TRƯNG (FEATURE ENGINEERING)", 1)

set_heading(doc, "3.1. Đặc trưng từ CICFlowMeter (67 features)", 2)
add_para(doc, "Dataset đã được trích xuất sẵn 67 features bởi CICFlowMeter — bao gồm:")
feat_groups = [
    ("Packet-level:", "Total Fwd/Bwd Packets, Flow Duration, Packet Length Mean/Std/Max, Packet Length Variance"),
    ("Flow-level:", "Flow Bytes/s, Flow Packets/s, Down/Up Ratio, Average Packet Size"),
    ("TCP flags:", "FIN Flag Count, SYN Flag Count, PSH Flag Count, ACK Flag Count"),
    ("Timing:", "Flow IAT Mean/Std/Min/Max, Active Mean/Std, Idle Mean/Std"),
    ("Port:", "Destination Port (quan trọng để lọc HTTP traffic)"),
]
for label, desc in feat_groups:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(desc)

set_heading(doc, "3.2. Đặc trưng cửa sổ (Windowed Features)", 2)
add_para(doc, "Trong pipeline đa luồng, gói tin được gom theo IP nguồn vào cửa sổ 60 giây. "
             "Từ mỗi cửa sổ, trích xuất:")
headers2 = ["Đặc trưng", "Công thức", "Ý nghĩa"]
rows2 = [
    ["upload_download_ratio", "fwd_bytes / max(bwd_bytes, 1)", "Exfil: ratio cao bất thường"],
    ["burst_count", "count(inter_arrival < 0.1s)", "Nhiều request liên tục → automated"],
    ["burst_ratio", "burst_count / total_intervals", "Tỷ lệ burst trên tổng"],
    ["unusual_port_ratio", "count(srv_port not in {80,443,8080,8443}) / n", "Port không phổ biến → suspicious"],
    ["inter_request_time_std", "std(inter_arrival_times)", "Exfil: std thấp (đều)"],
    ["request_rate", "count / duration(s)", "Tần suất request"],
    ["total_bytes", "sum(payload_len)", "Dung lượng transfer"],
    ["mean_payload_size", "mean(payload_lens)", "Kích thước payload trung bình"],
]
make_table(doc, headers2, rows2)

set_heading(doc, "3.3. Metric burst_exfil_score", 2)
add_para(doc, "Đây là metric chính do chúng tôi đề xuất, kết hợp 4 tín hiệu:")
score_items = [
    ("Upload ratio > 100 + bytes > 50KB:", "+0.40", "Tín hiệu mạnh nhất — Bot traffic upload 4.57x so với normal"),
    ("Burst count > 50:", "+0.20", "Nhiều request liên tục — automated behavior"),
    ("Unusual port ratio > 0.8:", "+0.20", "Port không phải HTTP/S standard"),
    ("Inter-request std < 0.05s:", "+0.20", "Machine-generated traffic — regular pattern"),
]
for feat, score, reason in score_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(feat).bold = True
    p.add_run(f" → {score}  ")
    p.add_run(f"({reason})")

add_para(doc, "Alert threshold: score > 0.7 → HIGH alert. Score = 1.0 → CRITICAL.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. PIPELINE ĐA LUỒNG
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "4. PIPELINE ĐA LUỒNG", 1)

set_heading(doc, "4.1. Kiến trúc 3 Threads", 2)

arch_text = """
┌──────────────────────────────────────────────────────────────────┐
│                        MAIN PROCESS                                 │
│  stop_event = threading.Event()                                  │
│                                                                  │
│  ┌──────────────┐    ┌──────────────┐    ┌──────────────────┐  │
│  │ THREAD 1     │    │ THREAD 2     │    │ THREAD 3         │  │
│  │ Packet       │    │ Feature      │    │ Inference +       │  │
│  │ Capture      │──→ │ Aggregation  │──→ │ Alert Logging     │  │
│  │              │    │              │    │                  │  │
│  │ Scapy sniff  │    │ 60s window  │    │ Load .pkl/.h5   │  │
│  │ BPF filter   │    │ per src IP   │    │ Predict + Alert  │  │
│  └──────────────┘    └──────────────┘    └──────────────────┘  │
│         │                   │                     │             │
│         ↓                   ↓                     ↓             │
│  packet_queue      feature_queue           exfil_detection.log │
│  (maxsize=50000)  (maxsize=10000)        console (RED)     │
└──────────────────────────────────────────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = "Courier New"
run.font.size = Pt(8)

set_heading(doc, "4.2. Chi tiết từng Thread", 2)

thread_details = [
    ("Thread 1 — Packet Capture (packet_capture.py)",
     ["Dùng Scapy sniff() đọc PCAP hoặc network interface.",
      "BPF filter: tcp (port 80 or port 443 or port 8080 or port 8443)",
      "Parse packet → dict {timestamp, src_ip, dst_ip, src_port, dst_port, payload_len, tcp_flags}",
      "Đẩy vào packet_queue (maxsize=50,000). Queue full → skip packet, không block.",
      "Hỗ trợ cả offline mode (PCAP) và live mode (sudo required)."]),
    ("Thread 2 — Feature Aggregation (feature_aggregator.py)",
     ["Buffer packets theo src_ip (dict: {src_ip: {packets, window_start}})",
      "Flush mỗi WINDOW_SIZE=60 giây hoặc khi shutdown.",
      "Direction detection: SYN/ACK heuristic + ephemeral port detection.",
      "Skip IPs có < MIN_PACKETS=3 packets trong window.",
      "Đẩy feature dict vào feature_queue (maxsize=10,000)."]),
    ("Thread 3 — Inference + Alert (model_inference.py)",
     ["Load scaler.pkl và model tại startup.",
      "Mỗi feature vector: compute burst_exfil_score → nếu > 0.7 → RED ALERT.",
      "CNN1D Final (tuned threshold) là model primary cho detection.",
      "Alert được log vào file (exfil_detection.log) và console.",
      "Graceful degradation nếu model load fail."]),
]
for title, items in thread_details:
    set_heading(doc, title, 3)
    for item in items:
        add_bullet(doc, item)

set_heading(doc, "4.3. Kết quả chạy Pipeline trên Friday-WorkingHours.pcap", 2)
add_para(doc, "Pipeline đã test thành công trên Friday-WorkingHours.pcap — "
             "chứa Friday DDoS attack (HOIC/LOIC) và PortScan traffic:")
headers3 = ["Thread", "Trạng thái", "Chi tiết"]
rows3 = [
    ["Thread 1 Capture", "✅ Hoạt động", "Đọc PCAP, parse packets, BPF filter đúng"],
    ["Thread 2 Aggregator", "✅ Hoạt động", "Buffer per IP, flush 60s, extract features"],
    ["Thread 3 Inference", "✅ Hoạt động", "Model load OK, burst_exfil_score firing alerts"],
    ["CNN1D Final", "✅ Loaded", "Predict với tuned threshold = 0.207"],
]
make_table(doc, headers3, rows3)

add_para(doc, "Alerts nổi bật phát hiện được:")
alert_headers = ["IP nguồn", "Severity", "Score", "Requests/window", "Bytes", "Pattern"]
alert_rows = [
    ["192.168.10.15", "CRITICAL", "1.000", "135,000+", "823 KB+", "Bot-like burst, 100% upload"],
    ["192.168.10.5/9/12/14/16", "HIGH", "0.800", "1,000–20,000", "70 KB–1.5 MB", "DDoS bot traffic"],
    ["72.247.71.7", "HIGH", "0.800", "207", "386 KB", "Massive upload burst"],
    ["91.189.88.149", "HIGH", "0.800", "139", "296 KB", "Exfiltration burst"],
]
make_table(doc, alert_headers, alert_rows)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. HUẤN LUYỆN MÔ HÌNH
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "5. HUẤN LUYỆN MÔ HÌNH", 1)

set_heading(doc, "5.1. Anomaly-based Models", 2)
add_para(doc, "Huấn luyện trên NORMAL traffic (BENIGN) để phát hiện outliers:")
add_para(doc, "Isolation Forest: n_estimators=200, contamination=0.10, max_samples=256")
add_para(doc, "One-Class SVM: kernel='rbf', gamma='scale', nu=0.05")

set_heading(doc, "5.2. Supervised Models (BiLSTM + CNN 1D)", 2)
add_para(doc, "Huấn luyện trên toàn bộ dataset đã gán nhãn (2,830,743 flows):")
add_para(doc, "Train/Test/Val: 70/15/15, stratified by label")

set_heading(doc, "5.3. Chiến lược xử lý Extreme Imbalance", 2)
add_para(doc, "Vấn đề gốc: với 0.078% exfil (1,543 exfil trên 1.98M train), "
             "training trực tiếp sẽ cho model có FPR rất cao (~45%).")
add_para(doc, "Giải pháp — Final Training Strategy (train_final.py):")

strategy_steps = [
    ("Step 1 — Subsample:", "Lấy 100,000 samples từ train, giữ nguyên tỷ lệ 1.6% exfil → đủ để model học patterns."),
    ("Step 2 — SMOTE 10%:", "Oversampling minority lên 10% của majority (~9,800 exfil, 98,000 normal) — "
     "chỉ 6× thay vì 128× như trước, giảm distortion."),
    ("Step 3 — Focal Loss α=0.50:", "Symmetric focal loss thay vì α=0.75 — ít bias hơn, cân bằng hơn."),
    ("Step 4 — class_weight={0:1, 1:5}:", "Moderate class weight thay vì 10× — giảm over-penalization."),
    ("Step 5 — Post-train Threshold Tuning:", "Sau khi train, quét ROC curve để tìm threshold tối ưu "
     "(FPR ≤ 5%, Recall ≥ 85%)."),
]
for label, desc in strategy_steps:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(label + " ")
    r1.bold = True
    p.add_run(desc)

add_para(doc, "CNN 1D: Conv1D(64→32) → GlobalAveragePooling → Dense(64) → Dense(1,sigmoid)")
add_para(doc, "BiLSTM: Bidirectional LSTM(64→32) → Dense(64) → BN → Dense(1,sigmoid)")
add_para(doc, "Optimizer: Adam, lr=0.001, epochs=30, batch_size=512, EarlyStopping(patience=7)")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. ĐÁNH GIÁ KẾT QUẢ
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "6. ĐÁNH GIÁ KẾT QUẢ", 1)

set_heading(doc, "6.1. Kết quả Final (sau Threshold Tuning) — Test Set 424,611 flows", 2)
add_para(doc, f"Tỷ lệ exfiltration trong test set: 0.074% (313 exfil / 424,611 total).")

headers4 = ["Mô hình", "AUC-ROC", "Recall", "FPR", "Precision", "F1", "Threshold"]
rows4 = [
    ["CNN 1D Final ⭐", "0.9971 ✅", "1.0000 ✅", "0.0245 ✅", "0.0292", "0.0567", "0.207"],
    ["BiLSTM Final", "0.9966 ✅", "1.0000 ✅", "0.0322 ✅", "0.0224", "0.0438", "0.167"],
]
make_table(doc, headers4, rows4)

add_para(doc, "✅ Mục tiêu đạt: AUC > 0.90, FPR < 5%, Recall ≥ 85%")

add_para(doc, "")
set_heading(doc, "6.2. Kết quả trước khi Threshold Tuning (threshold=0.5)", 2)
headers4b = ["Mô hình", "AUC-ROC", "Recall", "FPR", "Precision", "Note"]
rows4b = [
    ["CNN 1D", "0.9423 ✅", "1.0000", "0.4477 ❌", "0.0016", "FPR quá cao — cần threshold tuning"],
    ["BiLSTM", "0.9012 ✅", "1.0000", "0.4416 ❌", "0.0017", "FPR quá cao — cần threshold tuning"],
]
make_table(doc, headers4b, rows4b)

add_para(doc,
    "Nhận xét: Vấn đề không phải ở model discrimination (AUC tốt) mà ở threshold calibration. "
    "Sau khi quét ROC curve, threshold tối ưu là 0.21 (CNN1D) thay vì 0.5 — "
    "điều này làm giảm FPR từ ~0.45 xuống ~0.025 (giảm 18×).")

set_heading(doc, "6.3. Phân tích chi tiết", 2)
analyses = [
    ("CNN 1D Final (AUC=0.9971 — TỐT NHẤT):",
     "Đạt vượt target AUC>0.90 với margin lớn. FPR=2.45% đạt mục tiêu <5%. "
     "Recall=100% = phát hiện TẤT CẢ exfil attacks. Precision thấp (2.9%) là expected "
     "với extreme imbalance."),
    ("BiLSTM Final (AUC=0.9966):",
     "Vượt target AUC>0.90. Bidirectional giúp học temporal patterns trong feature sequences. "
     "FPR=3.22% — cũng đạt mục tiêu <5%."),
    ("Anomaly Models (AUC≈0.55):",
     "Kém vì Bot/BENIGN traffic giống nhau trong raw 67-feature space. "
     "Isolation Forest phù hợp cho zero-day detection hơn là structured attacks."),
    ("Root Cause FPR cao:",
     "SMOTE 128× oversampling + Focal Loss α=0.75 + class_weight=10× → "
     "model bị probability inflation. Mean probability = 0.44 trên dataset 99.9% normal. "
     "Giải pháp: Subsample + SMOTE 10% + Focal α=0.50 + Threshold Tuning."),
]
for label, desc in analyses:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label).bold = True
    p.add_run(f" {desc}")

set_heading(doc, "6.4. burst_exfil_score Threshold Analysis", 2)
thresh_headers = ["Threshold", "Mức độ", "Ưu điểm", "Nhược điểm"]
thresh_rows = [
    ["> 0.5", "Thấp", "Recall cao", "Nhiều false positives"],
    ["> 0.6", "Trung bình", "Cân bằng", "Có thể miss subtle exfil"],
    ["> 0.7 (default)", "Cao ✅", "Chấp nhận được", "Một số exfil nhỏ có thể miss"],
    ["> 0.8", "Rất cao", "Chỉ alert nghiêm trọng", "Có thể miss real exfil"],
]
make_table(doc, thresh_headers, thresh_rows)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. SO SÁNH ANOMALY VS SUPERVISED
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "7. SO SÁNH ANOMALY-BASED VÀ SUPERVISED", 1)

set_heading(doc, "7.1. Bảng so sánh chi tiết (với tuned thresholds)", 2)
comp_headers = ["Tiêu chí", "Isolation Forest", "One-Class SVM", "BiLSTM Final", "CNN 1D Final"]
comp_rows = [
    ["Cần label?", "Không", "Không", "Có", "Có"],
    ["AUC-ROC đạt được", "0.5277", "0.5546", "0.9966 ✅", "0.9971 ✅"],
    ["FPR (tuned)", "10.10%", "4.93%", "3.22% ✅", "2.45% ✅"],
    ["Recall", "3.83%", "4.47%", "100%", "100%"],
    ["Tốc độ Inference", "Rất nhanh", "Chậm", "Trung bình", "Nhanh"],
    ["Phát hiện Zero-day", "Tốt", "Tốt", "Kém", "Kém"],
    ["Real-time phù hợp?", "✅ Rất phù hợp", "❌ Không phù hợp", "✅ Phù hợp", "✅ Rất phù hợp"],
    ["Overfitting risk", "Thấp", "Thấp", "Trung bình", "Trung bình"],
]
make_table(doc, comp_headers, comp_rows)

set_heading(doc, "7.2. Khi nào nên dùng mô hình nào?", 2)
when_items = [
    ("CNN1D Final (khuyến nghị primary):",
     "Dùng làm primary detector — AUC=0.9971, FPR=2.45%, Recall=100%. "
     "Với tuned threshold=0.207, model phát hiện gần như tất cả exfil với số false alarms thấp."),
    ("BiLSTM Final (secondary):",
     "Dùng kết hợp với CNN1D để ensemble voting — tăng confidence khi cả hai đều alert."),
    ("Anomaly-based (Isolation Forest):",
     "Phù hợp khi không có label exfiltration, cần phát hiện zero-day attacks. "
     "Ưu tiên low FPR thay vì high recall."),
    ("burst_exfil_score:",
     "Dùng kết hợp với mọi model làm rule-based layer. "
     "Xử lý tốt automated exfiltration (high upload + burst pattern)."),
    ("Đề xuất deployment:",
     "CNN1D Final (primary) + burst_exfil_score (rule-based) + Isolation Forest (zero-day backup)."),
]
for label, desc in when_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label).bold = True
    p.add_run(f" {desc}")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. KẾT LUẬN VÀ ĐỀ XUẤT
# ══════════════════════════════════════════════════════════════════════════════
set_heading(doc, "8. KẾT LUẬN VÀ ĐỀ XUẤT", 1)

set_heading(doc, "8.1. Kết luận", 2)
conclusions = [
    "Pipeline đa luồng 3 threads hoạt động đúng đắn, xử lý PCAP real-time.",
    "CNN1D Final đạt AUC=0.9971, FPR=2.45% — vượt xa target AUC>0.90 và FPR<5%.",
    "BiLSTM Final đạt AUC=0.9966, FPR=3.22% — cũng vượt target.",
    "Threshold tuning là giải pháp quyết định: giảm FPR từ ~45% xuống ~2.5% (giảm 18×) "
     "mà không làm mất recall.",
    "burst_exfil_score phát hiện chính xác Bot traffic và DDoS patterns trên Friday PCAP.",
    "Anomaly models kém trên CICIDS2017 vì Bot traffic không đủ khác biệt trong raw feature space.",
    "Đề tài đáp ứng đầy đủ yêu cầu GVMH: phân tích HTTP traffic, trích đặc trưng cửa sổ, "
     "pipeline đa luồng, 4 mô hình ML, so sánh anomaly vs supervised, đề xuất metric riêng.",
]
for c in conclusions:
    add_bullet(doc, c)

set_heading(doc, "8.2. Hạn chế", 2)
limits = [
    "CICIDS2017 không có label exfiltration rõ ràng — dùng proxy (Bot, Infiltration) gây ảnh hưởng model quality.",
    "Anomaly models chỉ đạt AUC≈0.55 trên dataset này — cần dataset có label exfil thực sự.",
    "Pipeline chưa được test trên live network interface (cần sudo + network access).",
    "Precision thấp (~3%) là expected với extreme imbalance — cần production data để re-calibrate.",
]
for l in limits:
    add_bullet(doc, l)

set_heading(doc, "8.3. Đề xuất cải tiến", 2)
future = [
    "Tích hợp CNN1D Final vào pipeline thay cho Isolation Forest (sử dụng tuned threshold).",
    "Ensemble: kết hợp CNN1D + burst_exfil_score + Isolation Forest cho best-of-both-worlds.",
    "Tăng WINDOW_SIZE thêm parameter để linh hoạt theo traffic pattern.",
    "Thêm HTTP header inspection (User-Agent, Content-Type) để tăng detection accuracy.",
    "Thu thập production traffic data để re-train và calibrate thresholds theo real-world distribution.",
]
for f in future:
    add_bullet(doc, f)

set_heading(doc, "8.4. Tài liệu tham khảo", 2)
refs = [
    "CICIDS2017 Dataset — Canadian Institute for Cybersecurity: https://www.unb.ca/cic/datasets/ids-2017.html",
    "Isolation Forest — Liu et al. (2008): ACM TKDD",
    "Focal Loss — Lin et al. (2017): ICCV",
    "CICFlowMeter — IPCE 2018",
]
for r in refs:
    add_bullet(doc, r)

# Save
output_path = "/Users/nguyen_bao/Documents/AI_Project/Exfiltration/docs/bao_cao.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
