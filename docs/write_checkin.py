#!/usr/bin/env python3
"""
Generate docs/checkin_thay.docx — Mid-term check-in report sent to supervisor.
Run: python docs/write_checkin.py
Output: docs/checkin_thay.docx
"""

import datetime
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Output path ──────────────────────────────────────────────────────────────
HERE = Path(__file__).parent
OUT  = HERE / "checkin_thay.docx"


# ── Formatting helpers ────────────────────────────────────────────────────────
def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False, italic=False, indent_cm=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.75)
    if bold_prefix:
        p.add_run(bold_prefix + " ").bold = True
    p.add_run(text)
    return p


def shade_row(row, hex_color="D9E2F3"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)


def make_table(doc, headers, rows, header_hex="2E75B6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    shade_row(hdr_row, header_hex)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run  = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        if r_idx % 2 == 0:
            shade_row(row, "F2F2F2")
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_question_box(doc, q_num, question, context=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(f"❓ Câu hỏi {q_num}: {question}")
    run.bold = True
    run.font.size = Pt(11)
    if context:
        add_para(doc, context, italic=True, indent_cm=1.5)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

today = datetime.datetime.now().strftime("%d/%m/%Y")

# ── COVER PAGE ────────────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run("BÁO CÁO TIẾN ĐỘ VÀ XIN Ý KIẾN\nPHÁT HIỆN DATA EXFILTRATION QUA HTTP")
t.bold = True
t.font.size = Pt(18)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.add_run("Đồ án môn học — Giám sát an ninh mạng").bold = True

doc.add_paragraph()

for label, value in [
    ("Giáo viên hướng dẫn:", "ThS. Đàm Minh Linh  (linhdm@ptit.edu.vn)"),
    ("Học viện:", "Học viện Công nghệ Bưu Chính Viễn Thông — CS TP.HCM"),
    ("Ngày gửi:", today),
    ("Nhóm thực hiện:", "[Tên nhóm / MSSV]"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label + "  ").bold = True
    p.add_run(value)

doc.add_page_break()

# ── MỤC ĐÍCH ─────────────────────────────────────────────────────────────────
set_heading(doc, "I. MỤC ĐÍCH BÁO CÁO")
add_para(doc,
    "Báo cáo này trình bày tiến độ hiện tại của đồ án, lý do đưa ra các quyết định kỹ thuật "
    "(chọn mô hình, thiết kế hệ thống, phương pháp đánh giá), và xin ý kiến phản hồi của Thầy "
    "trước khi hoàn thiện báo cáo cuối kỳ và buổi demo trực tiếp.")

add_para(doc, "Các vấn đề cần xác nhận từ Thầy:", bold=True)
questions_summary = [
    "Chiến lược chọn model có phù hợp với yêu cầu đề bài không?",
    "Phương pháp gán nhãn heuristic có được chấp nhận không?",
    "Hình thức demo Thầy mong muốn (loopback / interface vật lý / slide minh họa)?",
    "Có cần bổ sung tính năng nào để hệ thống sát thực tế production hơn không?",
    "Thang điểm đánh giá — Thầy ưu tiên tiêu chí gì trong buổi bảo vệ?",
]
for q in questions_summary:
    add_bullet(doc, q)

doc.add_page_break()

# ── PHẦN 1: TỔNG QUAN HỆ THỐNG ───────────────────────────────────────────────
set_heading(doc, "II. TỔNG QUAN HỆ THỐNG ĐÃ TRIỂN KHAI")

set_heading(doc, "1. Kiến trúc 3 tầng — 3 luồng song song", 2)
add_para(doc,
    "Hệ thống được thiết kế theo mô hình Producer-Consumer với ba luồng (threads) chạy song song, "
    "kết nối qua hai hàng đợi thread-safe. Đây là bằng chứng xử lý đa luồng yêu cầu trong đề bài.")

arch_lines = [
    "  ┌──────────────┐   packet_queue    ┌──────────────┐   feature_queue   ┌─────────────────┐",
    "  │  Thread 1    │   (maxsize=50K)   │  Thread 2    │   (maxsize=10K)   │  Thread 3        │",
    "  │  Capture     │ ────────────────► │  Aggregator  │ ────────────────► │  Inference       │",
    "  │  (Scapy)     │                   │  (60s window)│                   │  (ML model)      │",
    "  └──────────────┘                   └──────────────┘                   └─────────────────┘",
    "        ↓                                                                        ↓",
    "  BPF tcp filter                                                    console RED + log file",
    "  port 80/443/8080                                                  + Telegram Bot alert",
]
p = doc.add_paragraph()
run = p.add_run("\n".join(arch_lines))
run.font.name = "Courier New"
run.font.size = Pt(8)

doc.add_paragraph()
set_heading(doc, "2. Chi tiết triển khai từng thread", 2)

threads = [
    ("Thread 1 — Packet Capture  (src/capture/packet_capture.py)",
     [
         "Dùng Scapy sniff() để đọc PCAP offline hoặc sniff live trên network interface.",
         "BPF filter: tcp and (port 80 or port 443 or port 8080 or port 8443).",
         "Output mỗi gói: {timestamp, src_ip, dst_ip, src_port, dst_port, payload_len, tcp_flags, pkt_len}.",
         "Khi packet_queue đầy → bỏ qua gói mới (không block), ghi cảnh báo.",
         "Hỗ trợ toggle OFFLINE_MODE = True/False trong config.py.",
     ]),
    ("Thread 2 — Feature Aggregation  (src/features/feature_aggregator.py)",
     [
         "Gom gói tin theo src_ip vào cửa sổ 60 giây (WINDOW_SIZE).",
         "Tự flush mỗi WINDOW_SIZE giây hoặc khi nhận stop_event.",
         "Tính các đặc trưng cửa sổ: upload_download_ratio, burst_count, burst_ratio, unusual_port_ratio, inter_request_time_std, request_rate, unique_destinations, v.v.",
         "Bỏ qua IP có ít hơn MIN_PACKETS=3 gói (lọc nhiễu).",
     ]),
    ("Thread 3 — Inference + Alert  (src/inference/model_inference.py)",
     [
         "Load scaler.pkl và model tại khởi động; warm-up bằng dummy predict.",
         "Mỗi feature vector: chuẩn hoá → dự đoán → tính burst_exfil_score.",
         "Alert nếu: prediction==1 HOẶC burst_exfil_score > burst_threshold (default=0.7).",
         "Ghi log vào exfil_detection.log (file) và in ra console màu đỏ.",
         "Tuỳ chọn: gửi Telegram alert qua Bot API (metadata only, không kèm payload).",
     ]),
]
for title, bullets in threads:
    set_heading(doc, title, 3)
    for b in bullets:
        add_bullet(doc, b)

doc.add_page_break()

# ── PHẦN 2: LÝ DO CHỌN MODEL ─────────────────────────────────────────────────
set_heading(doc, "III. LÝ DO CHỌN MÔ HÌNH")

set_heading(doc, "1. Đề xuất chọn đại diện theo từng trường phái", 2)
add_para(doc,
    "Đề bài liệt kê 4 mô hình: Isolation Forest, One-Class SVM (anomaly-based) "
    "và BiLSTM, CNN 1D (supervised). Nhóm đề xuất chọn MỘT đại diện mỗi trường phái "
    "để so sánh công bằng và tập trung phân tích:")

model_table_headers = ["Trường phái", "Model chọn", "Lý do loại bỏ đối thủ", "File"]
model_table_rows = [
    ["Anomaly-based",
     "Isolation Forest v2",
     "OCSVM: chậm ~10× trên 2.83M flows (không phù hợp real-time), kernel RBF tốn bộ nhớ với 67 features.",
     "isolation_forest_v2.pkl"],
    ["Supervised",
     "CNN 1D Final",
     "BiLSTM: model nặng hơn 9× (9.6MB vs 162KB), inference chậm hơn. CNN1D đạt AUC tương đương (0.9971 vs 0.9966).",
     "cnn1d_final.h5"],
]
make_table(doc, model_table_headers, model_table_rows)

doc.add_paragraph()
set_heading(doc, "2. Lý do chi tiết — Isolation Forest (anomaly)", 2)
if_reasons = [
    "Unsupervised — train chỉ trên BENIGN flows, không cần label.",
    "Inference O(log n) per sample → phù hợp pipeline real-time.",
    "Robust với high-dimensional data (67 features) hơn OCSVM.",
    "Dễ điều chỉnh contamination để ảnh hưởng FPR.",
    "Đã retrain v2 với 12 upload-related features và contamination=0.001 (≈ tỷ lệ exfil thực).",
]
for r in if_reasons:
    add_bullet(doc, r)

set_heading(doc, "3. Lý do chi tiết — CNN 1D (supervised)", 2)
cnn_reasons = [
    "1D Convolution trên feature vector học local pattern trong chuỗi đặc trưng.",
    "Lightweight: 162KB — có thể deploy edge device.",
    "Sau threshold tuning (0.207 thay vì 0.5): AUC=0.9971, Recall=100%, FPR=2.45%.",
    "Inference ~9 μs/flow trên CPU laptop — đáp ứng thời gian thực.",
    "Đã áp dụng: SMOTE 10%, Focal Loss, class_weight, EarlyStopping để xử lý extreme imbalance.",
]
for r in cnn_reasons:
    add_bullet(doc, r)

set_heading(doc, "4. Vì sao anomaly model kém hơn supervised trên CICIDS2017?", 2)
add_para(doc,
    "Đây là kết luận quan trọng của đề tài — Isolation Forest v2 chỉ đạt AUC=0.5463 "
    "(gần random) so với CNN1D AUC=0.9971. Nguyên nhân không phải lỗi thuật toán mà là "
    "giới hạn dữ liệu:")
anomaly_limits = [
    "CICIDS2017 không có label exfiltration thực. Nhóm dùng Bot + Infiltration làm proxy — label này noisy.",
    "Bot traffic trong CICIDS2017 có đặc trưng tương đối giống traffic bình thường khi nhìn từ 67 raw features.",
    "Trong không gian đặc trưng, pattern exfil chồng lấp với web upload hợp lệ (Google Drive, OneDrive).",
    "→ Kết luận khoa học hợp lệ: anomaly-based phù hợp phát hiện zero-day; supervised vượt trội khi có label.",
]
for l in anomaly_limits:
    add_bullet(doc, l)

doc.add_page_break()

# ── PHẦN 3: KẾT QUẢ ĐO LƯỜNG ─────────────────────────────────────────────────
set_heading(doc, "IV. KẾT QUẢ ĐÁNH GIÁ")

set_heading(doc, "1. Bảng so sánh trên test set (424,611 flows, 313 exfil)", 2)
add_para(doc, "Đánh giá trực tiếp bằng evaluate.py --compare trên cùng test set:", italic=True)

eval_headers = ["Model", "Loại", "AUC-ROC", "Recall", "FPR", "Precision", "F1", "Inf. time (μs)"]
eval_rows = [
    ["IsolationForest v1", "Anomaly", "0.5306", "3.83%", "5.00%", "0.06%", "0.0011", "—"],
    ["IsolationForest v2 ★", "Anomaly", "0.5463", "4.2%", "2.41%", "0.13%", "0.0025", "6.35"],
    ["CNN 1D Final ★", "Supervised", "0.9971", "100%", "2.45%", "2.92%", "0.0567", "8.95"],
    ["BiLSTM Final", "Supervised", "0.9966", "100%", "3.22%", "2.24%", "0.0438", "—"],
]
make_table(doc, eval_headers, eval_rows)

doc.add_paragraph()
add_para(doc,
    "★ = Đại diện được chọn. Δ AUC = 0.451 giữa supervised và anomaly. "
    "Tự động kết luận: 'Supervised vượt trội. Anomaly chỉ nên làm lớp bổ sung phát hiện zero-day.'",
    italic=True)

set_heading(doc, "2. Kết quả demo live (loopback lo0 — 30/04/2026)", 2)
add_para(doc, "Chạy pipeline live trên interface lo0 với attacker.py gửi 100 POST request 50KB:")

live_headers = ["Tiêu chí", "Kết quả", "Ghi chú"]
live_rows = [
    ["Thread 1 Capture", "✅ Hoạt động", "seen=2606 queued=2606 skipped=0"],
    ["Thread 2 Aggregator", "✅ Hoạt động", "windows=6 skipped=0"],
    ["Thread 3 Inference", "✅ Hoạt động", "processed=6 alerts=2"],
    ["Benign baseline (30 req, 200B, 1s delay)", "✅ Không alert", "score=0.200 < threshold 0.7"],
    ["Burst exfil (100 req, 50KB, 0.1s delay)", "✅ Alert đúng", "2 windows, HIGH severity trong 10–15s"],
    ["Upload ratio đo được", "496× – 497×", "3.67MB / 1.36MB upload vs 7KB download"],
    ["burst_exfil_score", "0.800", "Vượt ngưỡng 0.7"],
    ["Model score (runtime_window_rf.pkl)", "1.000", "RandomForestClassifier trên window features"],
    ["Telegram alert", "✅ Gửi thành công", "Metadata only, không chứa payload content"],
    ["Console duplicate print", "✅ Đã fix", "propagate=False trên exfil.alerts logger"],
]
make_table(doc, live_headers, live_rows)

doc.add_page_break()

# ── PHẦN 4: TÍNH NĂNG THỰC TẾ ────────────────────────────────────────────────
set_heading(doc, "V. CÁC TÍNH NĂNG HIỆN CÓ VÀ ĐỀ XUẤT THỰC TẾ")

set_heading(doc, "1. Tính năng đã triển khai", 2)
implemented = [
    "Live packet capture (Scapy) hoặc offline PCAP replay.",
    "BPF filter chỉ bắt HTTP/HTTPS traffic (port 80/443/8080/8443).",
    "Window-based feature aggregation per source IP (60s, configurable).",
    "Dual detection: ML model prediction + rule-based burst_exfil_score.",
    "Thread-safe logging (file + console colored) với graceful shutdown.",
    "Telegram Bot alert — chỉ gửi metadata, không chứa payload (privacy-compliant).",
    "CLI flags: --live/--offline, --window-size, --burst-threshold, --debug, --model.",
    "Offline mode fallback: replay PCAP khi không có network access.",
    "demo/attacker.py + demo/upload_server.py cho kịch bản demo có thể tái lập.",
    "Notebook so sánh: roc_comparison.png, comparison_table.png từ evaluate.py --compare.",
]
for i in implemented:
    add_bullet(doc, i)

set_heading(doc, "2. Tính năng có thể bổ sung nếu Thầy yêu cầu", 2)
optional_features = [
    ("HTTP Header Inspection:",
     "Phân tích User-Agent, Content-Type, Referer — phát hiện exfil cẩn thận hơn "
     "(ẩn data trong multipart/form-data). Đòi hỏi deep packet inspection, phức tạp hơn."),
    ("HTTPS/TLS Decryption:",
     "Giải mã TLS để inspect payload. Yêu cầu certificate pinning hoặc MITM proxy — "
     "phức tạp, có vấn đề pháp lý trong production."),
    ("Database Logging:",
     "Ghi alert vào SQLite/PostgreSQL thay vì chỉ file log — phù hợp production hơn."),
    ("Web Dashboard:",
     "Bảng điều khiển realtime (Flask + Chart.js) hiển thị traffic timeline và alert. "
     "Hữu ích cho demo trực quan nhưng không cốt lõi."),
    ("Self-captured PCAP scenarios:",
     "Chụp PCAP thực từ máy lab với kịch bản exfil thật (curl, nc, python) — "
     "tăng điểm thực tiễn theo AGENTS.md. Hiện đã có demo_exfil_local.pcap cơ bản."),
    ("Anomaly ensemble:",
     "Kết hợp IsolationForest (unsupervised) + CNN1D (supervised) theo voting — "
     "giảm FP khi một model không chắc chắn."),
]
for feat, desc in optional_features:
    add_bullet(doc, desc, bold_prefix=feat)

doc.add_page_break()

# ── PHẦN 5: CÂU HỎI XIN Ý KIẾN ──────────────────────────────────────────────
set_heading(doc, "VI. CÂU HỎI XIN Ý KIẾN THẦY")

add_para(doc,
    "Nhóm rất mong nhận phản hồi từ Thầy về 5 vấn đề dưới đây trước khi hoàn thiện "
    "báo cáo cuối kỳ và chuẩn bị demo:")

doc.add_paragraph()
add_question_box(
    doc, 1,
    "Chiến lược chọn model — 2 đại diện có đúng ý Thầy không?",
    'Nhóm chọn Isolation Forest v2 (anomaly) và CNN 1D Final (supervised) làm 2 đại diện '
    'để so sánh công bằng thay vì chạy cả 4 model. Lý do: OCSVM quá chậm cho real-time, '
    'BiLSTM nặng hơn CNN1D mà kết quả tương đương. '
    'Nếu Thầy yêu cầu trình bày đủ cả 4 model (chỉ bảng số, không cần demo live), nhóm sẽ bổ sung.'
)

doc.add_paragraph()
add_question_box(
    doc, 2,
    "Gán nhãn heuristic — có chấp nhận được trong báo cáo không?",
    'CICIDS2017 không có label exfiltration thực sự. Nhóm dùng Bot + Infiltration traffic làm proxy '
    '(2,204 flows / 2.83M total) và thêm upload-ratio heuristic. Việc này được giải thích rõ trong '
    'báo cáo và đây là thực tiễn phổ biến trong nghiên cứu IDS. '
    'Nếu Thầy yêu cầu ground-truth label thực, nhóm cần tự capture PCAP exfil trên máy lab.'
)

doc.add_paragraph()
add_question_box(
    doc, 3,
    "Hình thức demo — Thầy mong muốn hình thức nào?",
    'Hiện tại nhóm đã chạy thử thành công trên lo0 (loopback 127.0.0.1). '
    'Nhóm muốn xác nhận:\n'
    '  (a) Demo trên loopback lo0 có được tính là "real-time detection" hợp lệ?\n'
    '  (b) Hay Thầy yêu cầu capture trên interface vật lý Wi-Fi/Ethernet?\n'
    '  (c) Hay chỉ cần demo offline với PCAP là đủ?\n'
    'Nhóm có thể làm cả 3, nhưng cần ưu tiên để chuẩn bị tốt nhất.'
)

doc.add_paragraph()
add_question_box(
    doc, 4,
    "Tính năng bổ sung — Thầy có yêu cầu thêm gì để sát thực tế production không?",
    'Các tính năng nhóm có thể thêm nếu cần (mỗi tính năng ~2–4 giờ):\n'
    '  • Database logging thay vì file log\n'
    '  • Web dashboard realtime (Flask)\n'
    '  • HTTP header inspection (User-Agent, Content-Type)\n'
    '  • Thêm kịch bản tự capture PCAP exfil từ máy lab\n'
    'Nhóm muốn biết Thầy ưu tiên hướng nào để tập trung.'
)

doc.add_paragraph()
add_question_box(
    doc, 5,
    "Tiêu chí đánh giá — Thầy ưu tiên điểm gì trong buổi bảo vệ?",
    'Để chuẩn bị tốt, nhóm muốn biết trọng số của các hạng mục:\n'
    '  • Kết quả mô hình (AUC, FPR, Recall)?\n'
    '  • Demo live pipeline chạy được?\n'
    '  • Chất lượng báo cáo / phân tích?\n'
    '  • Tính năng bổ sung sáng tạo?\n'
    '  • Khả năng giải thích quyết định kỹ thuật khi được hỏi?'
)

doc.add_paragraph()
doc.add_paragraph()
add_para(doc,
    "Nhóm rất mong nhận phản hồi sớm từ Thầy. Xin trân trọng cảm ơn!",
    bold=True)

doc.add_paragraph()
add_para(doc, f"Kính gửi,", italic=True)
add_para(doc, "[Tên nhóm / Đại diện nhóm]", italic=True)
add_para(doc, today, italic=True)

doc.add_page_break()

# ── PHẦN 6: PHỤ LỤC ─────────────────────────────────────────────────────────
set_heading(doc, "PHỤ LỤC — TÌNH TRẠNG TIẾN ĐỘ")

progress_headers = ["Hạng mục", "Tình trạng", "Ghi chú"]
progress_rows = [
    ["Phase 1: Setup + EDA", "✅ 100%", "notebooks/01_EDA.ipynb"],
    ["Phase 2: Feature Engineering", "✅ 100%", "2.83M flows, scaler đã fit"],
    ["Phase 3: Multi-thread Pipeline", "✅ 100%", "Test xong offline + live lo0"],
    ["Phase 4: Model Training", "✅ 100%", "IF_v2 + CNN1D Final + BiLSTM Final"],
    ["Phase 5: Evaluation", "✅ 100%", "evaluate.py --compare, PNG/JSON/MD"],
    ["Phase 6: Báo cáo + Demo", "⏳ 80%", "Chờ phản hồi Thầy để finalize"],
    ["Telegram Alert", "✅ Hoạt động", "Alert gửi được, metadata-only"],
    ["Live Demo lo0", "✅ Thành công", "alerts=2, no console duplicate"],
]
make_table(doc, progress_headers, progress_rows)

doc.add_paragraph()
set_heading(doc, "Cấu trúc thư mục chính", 2)
tree_lines = [
    "Exfiltration/",
    "├── src/",
    "│   ├── capture/        ← Thread 1: Scapy packet capture",
    "│   ├── features/       ← Thread 2: window aggregation + burst_exfil_score",
    "│   ├── inference/      ← Thread 3: ML predict + alert logger (Telegram)",
    "│   ├── train/          ← preprocess.py, train_dl.py, evaluate.py --compare",
    "│   └── pipeline.py     ← main entry point (--live/--offline, --debug)",
    "├── models/",
    "│   ├── isolation_forest_v2.pkl  ← Anomaly đại diện (retrain HTTP-only features)",
    "│   ├── cnn1d_final.h5           ← Supervised đại diện (AUC=0.9971)",
    "│   └── runtime_window_rf.pkl    ← Model chạy trong live pipeline",
    "├── data/processed/",
    "│   ├── train.csv / test.csv / val.csv",
    "│   └── evaluation_results.json",
    "├── notebooks/",
    "│   ├── 03_Model_Comparison.ipynb",
    "│   ├── roc_comparison.png",
    "│   └── comparison_table.png",
    "└── demo/",
    "    ├── attacker.py     ← Kịch bản burst exfil",
    "    └── upload_server.py",
]
p = doc.add_paragraph()
run = p.add_run("\n".join(tree_lines))
run.font.name = "Courier New"
run.font.size = Pt(8.5)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"✅  Saved: {OUT}")
